"""
Read each document under the corpus folder, normalize text, write plain text to
``processed_data`` using the same base name as the source file (``.txt``).

Pipeline (per file): lowercase → strip URLs / punctuation / digits → tokenize →
remove English stopwords → drop tokens shorter than mean/3 or longer than 3× mean
(mean = average character length after stopword removal) → lemmatize (NLTK).

Requires: ``pip install python-docx nltk`` and NLTK corpora (downloaded on first run).
"""

from __future__ import annotations

import argparse
import re
import string
import sys
from pathlib import Path

URL_RE = re.compile(
    r"(https?://[^\s]+)|(www\.[^\s]+)",
    re.IGNORECASE,
)


def parse_args() -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[2]
    default_in = _resolve_input_dir(repo_root)
    default_out = repo_root / "processed_data"

    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument(
        "--input-dir",
        type=Path,
        default=default_in,
        help=f"source folder (default: {default_in})",
    )
    p.add_argument(
        "--output-dir",
        type=Path,
        default=default_out,
        help=f"output folder (default: {default_out})",
    )
    p.add_argument(
        "--limit",
        type=int,
        default=0,
        help="process at most N files (0 = all)",
    )
    return p.parse_args()


def _resolve_input_dir(repo_root: Path) -> Path:
    for name in ("original_ data", "original_data"):
        p = repo_root / name
        if p.is_dir():
            return p
    return repo_root / "original_ data"


def ensure_nltk_data() -> None:
    import nltk

    checks: list[tuple[str, str]] = [
        ("stopwords", "corpora/stopwords"),
        ("wordnet", "corpora/wordnet"),
        ("omw-1.4", "corpora/omw-1.4"),
        ("averaged_perceptron_tagger_eng", "taggers/averaged_perceptron_tagger_eng"),
    ]
    for pkg, rel in checks:
        try:
            nltk.data.find(rel)
        except LookupError:
            nltk.download(pkg, quiet=True)


def extract_text_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(t)

    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".docx":
        return extract_text_docx(path)
    if suf == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"unsupported extension: {path.suffix}")


def preprocess_text(raw: str) -> str:
    from nltk import pos_tag
    from nltk.corpus import stopwords, wordnet
    from nltk.stem import WordNetLemmatizer

    text = raw.lower()
    text = URL_RE.sub(" ", text)
    text = text.translate(str.maketrans("", "", string.punctuation))
    text = re.sub(r"\d+", " ", text)
    tokens = [t for t in text.split() if t]

    sw = set(stopwords.words("english"))
    tokens = [t for t in tokens if t not in sw]

    if not tokens:
        return ""

    lengths = [len(t) for t in tokens]
    mean_len = sum(lengths) / len(lengths)
    lo = mean_len / 3.0
    hi = 3.0 * mean_len
    filtered: list[str] = []
    for t in tokens:
        L = len(t)
        if L < lo or L > hi:
            continue
        filtered.append(t)
    if not filtered:
        return ""

    def wn_pos(tag: str) -> str:
        if tag.startswith("J"):
            return wordnet.ADJ
        if tag.startswith("V"):
            return wordnet.VERB
        if tag.startswith("N"):
            return wordnet.NOUN
        if tag.startswith("R"):
            return wordnet.ADV
        return wordnet.NOUN

    lem = WordNetLemmatizer()
    tagged = pos_tag(filtered)
    lemmas = [lem.lemmatize(w, wn_pos(p)) for w, p in tagged]
    return " ".join(lemmas)


def main() -> int:
    args = parse_args()
    in_dir: Path = args.input_dir.resolve()
    out_dir: Path = args.output_dir.resolve()

    if not in_dir.is_dir():
        print(f"error: input not a directory: {in_dir}", file=sys.stderr)
        return 1

    out_dir.mkdir(parents=True, exist_ok=True)

    ensure_nltk_data()

    files = sorted(
        p
        for p in in_dir.iterdir()
        if p.is_file() and not p.name.startswith("~$")
    )
    if args.limit > 0:
        files = files[: args.limit]

    ok = 0
    for path in files:
        try:
            raw = extract_text(path)
        except ValueError as e:
            print(f"skip {path.name!r}: {e}", file=sys.stderr)
            continue
        except Exception as e:
            print(f"skip {path.name!r}: read error: {e}", file=sys.stderr)
            continue

        try:
            processed = preprocess_text(raw)
        except Exception as e:
            print(f"skip {path.name!r}: preprocess error: {e}", file=sys.stderr)
            continue

        out_name = f"{path.stem}.txt"
        out_path = out_dir / out_name
        out_path.write_text(processed + ("\n" if processed else ""), encoding="utf-8")
        ok += 1
        print(out_name)

    print(f"wrote {ok} file(s) to {out_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
