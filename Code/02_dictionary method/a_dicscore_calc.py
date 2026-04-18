"""
Calculate dictionary scores for each speech and append them to the Excel table.

For each dictionary CSV under ``dict/``, the script computes:

    dictionary_score = matched_term_occurrences / total_words * 100

The score is written as a new column into ``processed_data/processed_data_output.xlsx``.

Special handling:
- most dictionaries are calculated from the existing ``Text`` column
- ``we_gt`` and ``they_gt`` are calculated from the original raw speech documents in
  ``original_ data`` so they are not distorted by stopword removal

Dictionary CSV format:
- one column named ``term``
- entries may be single words or multi-word phrases
- entries ending with ``*`` are matched by prefix (fuzzy stem match)
"""

from __future__ import annotations

import argparse
import csv
import re
import sys
from dataclasses import dataclass
from pathlib import Path

TEXT_COL = "Text"
TOKEN_RE = re.compile(r"\S+")
RAW_TOKEN_RE = re.compile(r"[A-Za-z']+")
RAW_SOURCE_DICTS = {"we_gt", "they_gt"}


@dataclass(frozen=True)
class DictPattern:
    name: str
    parts: tuple[str, ...]


@dataclass(frozen=True)
class SourceDoc:
    path: Path
    speaker: str
    year: int
    month: str


def parse_args() -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[2]
    default_dict_dir = repo_root / "dict"
    default_xlsx = repo_root / "processed_data" / "processed_data_output.xlsx"
    default_original_dir = repo_root / "original_ data"

    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument(
        "--dict-dir",
        type=Path,
        default=default_dict_dir,
        help=f"directory containing dictionary CSV files (default: {default_dict_dir})",
    )
    p.add_argument(
        "--xlsx",
        type=Path,
        default=default_xlsx,
        help=f"Excel file to update in place (default: {default_xlsx})",
    )
    p.add_argument(
        "--original-dir",
        type=Path,
        default=default_original_dir,
        help=f"directory containing original speech files (default: {default_original_dir})",
    )
    return p.parse_args()


def normalize_text(text: str) -> list[str]:
    return TOKEN_RE.findall((text or "").lower())


def normalize_raw_text(text: str) -> list[str]:
    return RAW_TOKEN_RE.findall((text or "").lower())


def normalize_term(term: str) -> tuple[str, ...]:
    return tuple(part for part in term.strip().lower().split() if part)


def token_matches(pattern_part: str, token: str) -> bool:
    if pattern_part.endswith("*"):
        stem = pattern_part[:-1]
        return token.startswith(stem)
    return token == pattern_part


def phrase_matches(tokens: list[str], start: int, pattern: DictPattern) -> bool:
    if start + len(pattern.parts) > len(tokens):
        return False
    for offset, pattern_part in enumerate(pattern.parts):
        if not token_matches(pattern_part, tokens[start + offset]):
            return False
    return True


def count_matches(tokens: list[str], patterns: list[DictPattern]) -> int:
    total = 0
    for i in range(len(tokens)):
        for pattern in patterns:
            if phrase_matches(tokens, i, pattern):
                total += 1
    return total


def load_dictionary(path: Path) -> list[DictPattern]:
    patterns: list[DictPattern] = []
    seen: set[tuple[str, ...]] = set()

    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        if "term" not in (reader.fieldnames or []):
            raise ValueError(f"dictionary missing 'term' column: {path}")
        for row in reader:
            raw = (row.get("term") or "").strip()
            if not raw:
                continue
            parts = normalize_term(raw)
            if not parts or parts in seen:
                continue
            seen.add(parts)
            patterns.append(DictPattern(name=raw, parts=parts))

    if not patterns:
        raise ValueError(f"dictionary has no terms: {path}")
    return patterns


def column_name_from_dict(path: Path) -> str:
    return f"{path.stem}_score"


def parse_original_doc(path: Path) -> SourceDoc | None:
    if path.suffix.lower() not in {".docx", ".txt"}:
        return None
    parts = path.stem.split()
    if len(parts) != 3:
        return None
    year_letter, speaker, month = parts
    if len(year_letter) != 5 or not year_letter[:4].isdigit():
        return None
    return SourceDoc(
        path=path,
        speaker=speaker,
        year=int(year_letter[:4]),
        month=month,
    )


def extract_text_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(t)

    return "\n".join(parts)


def extract_source_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return extract_text_docx(path)
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"unsupported extension: {path.suffix}")


def build_original_doc_index(original_dir: Path) -> dict[tuple[str, int, str], Path]:
    index: dict[tuple[str, int, str], Path] = {}
    for path in sorted(original_dir.iterdir()):
        if not path.is_file() or path.name.startswith("~$"):
            continue
        parsed = parse_original_doc(path)
        if parsed is None:
            continue
        key = (parsed.speaker.lower(), parsed.year, parsed.month)
        if key not in index:
            index[key] = parsed.path
    return index


def main() -> int:
    args = parse_args()
    dict_dir = args.dict_dir.resolve()
    xlsx_path = args.xlsx.resolve()
    original_dir = args.original_dir.resolve()

    if not dict_dir.is_dir():
        print(f"error: dictionary directory not found: {dict_dir}", file=sys.stderr)
        return 1
    if not xlsx_path.is_file():
        print(f"error: Excel file not found: {xlsx_path}", file=sys.stderr)
        return 1
    if not original_dir.is_dir():
        print(f"error: original speech directory not found: {original_dir}", file=sys.stderr)
        return 1

    try:
        import pandas as pd
    except ImportError:
        print("error: need pandas and openpyxl (pip install pandas openpyxl)", file=sys.stderr)
        return 1

    df = pd.read_excel(xlsx_path)
    required_cols = {TEXT_COL, "Speaker", "Year", "Month"}
    missing_cols = required_cols.difference(df.columns)
    if missing_cols:
        print(f"error: Excel file missing required columns: {sorted(missing_cols)}", file=sys.stderr)
        return 1

    dict_files = sorted(
        p for p in dict_dir.glob("*.csv") if p.is_file() and not p.name.startswith("~$")
    )
    if not dict_files:
        print(f"error: no dictionary CSV files found in {dict_dir}", file=sys.stderr)
        return 1

    token_cache = [normalize_text(str(text)) for text in df[TEXT_COL].fillna("")]
    original_index = build_original_doc_index(original_dir)
    raw_token_cache: list[list[str] | None] = []
    for _, row in df.iterrows():
        key = (str(row["Speaker"]).strip().lower(), int(row["Year"]), str(row["Month"]).strip())
        path = original_index.get(key)
        if path is None:
            raw_token_cache.append(None)
            continue
        raw_text = extract_source_text(path)
        raw_token_cache.append(normalize_raw_text(raw_text))

    for dict_path in dict_files:
        patterns = load_dictionary(dict_path)
        col = column_name_from_dict(dict_path)
        scores: list[float] = []
        use_raw_source = dict_path.stem in RAW_SOURCE_DICTS
        for i, tokens in enumerate(token_cache):
            current_tokens = raw_token_cache[i] if use_raw_source else tokens
            if current_tokens is None:
                scores.append(0.0)
                continue
            if not current_tokens:
                scores.append(0.0)
                continue
            matched = count_matches(current_tokens, patterns)
            score = matched / len(current_tokens) * 100.0
            scores.append(score)
        df[col] = scores
        print(f"added column: {col}")

    df.to_excel(xlsx_path, index=False, engine="openpyxl")
    print(f"updated Excel file: {xlsx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
